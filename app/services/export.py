"""Export module: JSON, Markdown, DOCX."""
import json
import logging
from pathlib import Path
from sqlalchemy.orm import Session
from app.models.topic import Topic
from app.models.pipeline import (
    SynthesisResult, TaxonomyCandidate, GapRecord,
    DraftSection, ReviewReport
)

logger = logging.getLogger(__name__)


def build_export_bundle(topic: Topic, db: Session) -> dict:
    """Assemble all results into a single export dict."""
    synthesis = (
        db.query(SynthesisResult).filter_by(topic_id=topic.id)
        .order_by(SynthesisResult.created_at.desc()).first()
    )
    taxonomy = (
        db.query(TaxonomyCandidate).filter_by(topic_id=topic.id)
        .order_by(TaxonomyCandidate.created_at.desc()).first()
    )
    gaps = db.query(GapRecord).filter_by(topic_id=topic.id).all()
    drafts = db.query(DraftSection).filter_by(topic_id=topic.id).all()
    review = (
        db.query(ReviewReport).filter_by(topic_id=topic.id)
        .order_by(ReviewReport.created_at.desc()).first()
    )

    papers_export = []
    for p in topic.papers:
        entry = {
            "id": p.id, "title": p.title, "authors": p.authors,
            "year": p.year, "venue": p.venue, "url": p.url,
            "decision": {
                "label": p.decision.label,
                "score": p.decision.relevance_score,
                "reason": p.decision.reason,
            } if p.decision else None,
            "extraction": {
                "method_type": p.extraction.method_type,
                "setting": p.extraction.setting,
                "datasets": p.extraction.datasets,
                "strengths": p.extraction.strengths,
                "limitations": p.extraction.limitations,
            } if p.extraction else None,
        }
        papers_export.append(entry)

    return {
        "topic": {"id": topic.id, "title": topic.title, "description": topic.description},
        "query_plan": {
            "bundles": [
                {"label": b.label, "query_text": b.query_text, "source": b.source}
                for b in (topic.query_plans[-1].bundles if topic.query_plans else [])
            ]
        } if topic.query_plans else None,
        "papers": papers_export,
        "synthesis": {
            "comparison_table": synthesis.comparison_table if synthesis else None,
            "recurring_patterns": synthesis.recurring_patterns if synthesis else None,
            "contradictions": synthesis.contradictions if synthesis else None,
            "method_clusters": synthesis.method_clusters if synthesis else None,
        } if synthesis else None,
        "taxonomy": {
            "dimensions": taxonomy.dimensions if taxonomy else None,
            "paper_mapping": taxonomy.paper_mapping if taxonomy else None,
            "explanation": taxonomy.explanation if taxonomy else None,
        } if taxonomy else None,
        "gaps": [
            {"gap_type": g.gap_type, "title": g.gap_type, "description": g.description,
             "priority": g.priority, "evidence_ids": g.evidence_paper_ids}
            for g in gaps
        ],
        "draft_sections": [
            {"section": d.section_name, "version": d.version, "content": d.content}
            for d in drafts
        ],
        "review": {
            "major_weaknesses": review.major_weaknesses,
            "minor_issues": review.minor_issues,
            "revision_priorities": review.revision_priorities,
            "overall_score": review.overall_score,
        } if review else None,
    }


def export_markdown(bundle: dict) -> str:
    """Convert export bundle to Markdown."""
    lines = [f"# Survey: {bundle['topic']['title']}\n"]

    lines.append("## Paper Library\n")
    for p in bundle["papers"]:
        label = p["decision"]["label"] if p["decision"] else "unscreened"
        lines.append(f"- [{label}] **{p['title']}** ({p['year']}) - {p.get('venue','')}")
    lines.append("")

    if bundle.get("taxonomy"):
        lines.append("## Taxonomy\n")
        lines.append(bundle["taxonomy"].get("explanation", ""))
        lines.append("")

    if bundle.get("gaps"):
        lines.append("## Research Gaps\n")
        for g in bundle["gaps"]:
            lines.append(f"- **[{g['priority']}] {g['type']}**: {g['description']}")
        lines.append("")

    if bundle.get("draft_sections"):
        lines.append("## Draft Sections\n")
        seen: set[str] = set()
        for d in bundle["draft_sections"]:
            if d["section"] not in seen:
                seen.add(d["section"])
                lines.append(f"### {d['section'].replace('_', ' ').title()}\n")
                lines.append(d["content"])
                lines.append("")

    if bundle.get("review"):
        lines.append("## Reviewer Feedback\n")
        r = bundle["review"]
        lines.append(f"**Overall Score**: {r.get('overall_score', 'N/A')}\n")
        lines.append(f"**Major Weaknesses**:\n{r.get('major_weaknesses', '')}\n")
        lines.append(f"**Revision Priorities**:\n{r.get('revision_priorities', '')}\n")

    return "\n".join(lines)


def export_docx(bundle: dict, output_path: str) -> str:
    """Export to DOCX format."""
    from docx import Document
    doc = Document()
    doc.add_heading(f"Survey: {bundle['topic']['title']}", 0)

    if bundle.get("draft_sections"):
        seen: set[str] = set()
        for d in bundle["draft_sections"]:
            if d["section"] not in seen:
                seen.add(d["section"])
                doc.add_heading(d["section"].replace("_", " ").title(), level=1)
                doc.add_paragraph(d["content"])

    if bundle.get("review"):
        doc.add_heading("Reviewer Feedback", level=1)
        r = bundle["review"]
        doc.add_paragraph(f"Score: {r.get('overall_score', 'N/A')}")
        doc.add_paragraph(f"Major Weaknesses:\n{r.get('major_weaknesses', '')}")

    doc.save(output_path)
    return output_path
